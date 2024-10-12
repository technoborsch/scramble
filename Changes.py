import os
import docx
from copy import copy
from typing import Any, Optional

from config import DOC_SIZES_MAP, FORMAT_INFO, DEFAULT_GEOMETRY
from tools import unzip_page_numbers, get_latest_changes


class Changes:
    """
    Класс, описывающий изменения, которые будут внесены в извещение об изменении.
    """
    def __init__(self, folder_path: str) -> None:
        self._full_changes = self._extract(folder_path)
        self._changes = get_latest_changes(self._full_changes)

    def _extract(self, folder_path: str) -> dict[Any, dict]:
        """
        Извлекает изменения, описанные в файлах *AB внутри указанной папки.
        :param folder_path: Путь до рабочей папки с извещением об изменении
        :return: Словарь с описанием изменений
        """
        doc_paths = map(lambda x: os.path.join(folder_path, x), self._find_list_of_documents_of_the_set(folder_path))
        changes = {}
        for doc_path in doc_paths:
            doc = docx.Document(doc_path)
            set_code = doc.tables[0].rows[1].cells[0].text.split("-")[0].replace("\n", "")
            total_sheets = 0
            for table in doc.tables:
                for row in table.rows[1:-2]:
                    row = row.cells
                    revision = row[2].text.lower()
                    split_revision = revision.split()
                    number_of_sheets = split_revision[0].split("/")[1].split(".")[1]
                    if "изм" in revision:
                        if set_code not in changes.keys():
                            changes[set_code] = {}
                        doc_code = row[0].text.replace("\n", "")
                        if doc_code not in changes[set_code].keys():
                            doc_ru_name, doc_eng_name = row[1].text.split("/")
                            if "MFS" in doc_code.split("-")[1]:
                                doc_ru_name += row[1].text.split()[-1]
                            revision_number = split_revision[0].split("/")[0]
                            set_position = split_revision[0].split("/")[1].split(".")[0]
                            changes[set_code][doc_code] = {
                                "doc_ru_name": doc_ru_name.strip(),
                                "doc_eng_name": doc_eng_name.strip(),
                                "revision": revision_number,
                                "number_of_sheets": number_of_sheets,
                                "set_start_page": total_sheets,
                                "set_position": set_position,
                                "changes": [],
                                "has_archive_number": self._resolve_archive_number(doc_code),
                                "page_size": self._resolve_page_size(doc_code),
                                "geometry": {},
                            }
                        extracted_changes = self._extract_changed_sheets(revision, number_of_sheets)
                        changes[set_code][doc_code]["changes"] = extracted_changes
                        changes[set_code][doc_code]["geometry"] = self._fill_geometry(doc_code, extracted_changes)
                    total_sheets += int(number_of_sheets)
            for set_code, set_changes in changes.items():
                changes[set_code] = dict(sorted(set_changes.items(), key=lambda item: int(item[1]["set_position"])))
        return dict(sorted(changes.items(), key=lambda item: item[0]))

    @staticmethod
    def _extract_changed_sheets(revision: str, number_of_sheets: str) -> list[dict]:
        """
        Извлекает изменения из текста описания изменений.
        :param revision: Текст изменений
        :param number_of_sheets: число страниц документа
        :return: список изменений
        """
        changes = [change.strip(". ") for change in revision.split("изм")]
        changes_list = []
        for change in changes[1:]:
            split_change = change.split()
            change_number = int(split_change[0])
            if len(split_change) == 1:
                changes_list.append({
                    "change_number": change_number,
                    "change_type": "patch",
                    "change_description_ru": "",
                    "change_description_en": "",
                    "pages": [num + 1 for num in range(0, int(number_of_sheets))],
                    "sections_number": []
                })
            content = split_change[1:]
            content_string = " ".join(content)
            items = content_string.split(")")
            for item in items:
                if item:
                    pages, type_ = item.split("(")
                    if pages:
                        pages = unzip_page_numbers(pages.strip(", "))
                    else:
                        pages = [num + 1 for num in range(0, int(number_of_sheets))]
                    if "зам" in type_:
                        changes_list.append({
                            "change_number": change_number,
                            "change_type": "replace",
                            "change_description_ru": "",
                            "change_description_en": "",
                            "pages": pages
                        })
                    elif "нов" in type_:
                        changes_list.append({
                            "change_number": change_number,
                            "change_type": "new",
                            "change_description_ru": "",
                            "change_description_en": "",
                            "pages": pages
                        })
                    elif "анн" in type_:
                        changes_list.append({
                            "change_number": change_number,
                            "change_type": "cancel",
                            "change_description_ru": "",
                            "change_description_en": "",
                            "pages": pages
                        })
                    elif "уч" in type_ or "изм" in type_:
                        changes_list.append({
                            "change_number": change_number,
                            "change_type": "patch",
                            "change_description_ru": "",
                            "change_description_en": "",
                            "pages": pages,
                            "sections_number": []
                        })
        return changes_list

    @staticmethod
    def _find_list_of_documents_of_the_set(folder_path: str) -> list[str]:
        """
        Метод ищет файлы AB в указанной папке
        :param folder_path: папка с рабочими файлами извещения об изменении
        :return:
        """
        list_of_paths = []
        for filename in os.listdir(folder_path):
            split_filename = filename.split("-")
            if len(split_filename) > 1 and "AB" in split_filename[1][1:] and filename.endswith("docx"):
                list_of_paths.append(filename)
        return list_of_paths

    @staticmethod
    def _resolve_archive_number(doc_code: str) -> bool:
        """
        Метод возвращает ответ - нужен ли документу с таким кодом архивный номер
        :param doc_code: Код документа
        :return: Нужен ли документу архивный номер
        """
        doc_letters = doc_code.split("-")[-1][:3]
        if doc_letters in DOC_SIZES_MAP.keys():
            return FORMAT_INFO[DOC_SIZES_MAP[doc_letters]]["archive_stamp"]
        else:
            return False

    @staticmethod
    def _resolve_page_size(doc_code: str) -> Optional[None, str]:
        """
        Возвращает размер формата документа, если знает про него
        :param doc_code: Код документа
        :return: Формат документа как строку или None
        """
        doc_letters = doc_code.split("-")[-1][:3]
        if doc_letters in DOC_SIZES_MAP.keys():
            return DOC_SIZES_MAP[doc_letters]
        else:
            return None

    @staticmethod
    def _fill_geometry(doc_code: str, changes: list[dict]) -> list[tuple]:
        """
        Метод заполняет геометрию документа значениями из файла конфигураций
        :param doc_code: Код документа
        :param changes: Лист с изменениями документа
        :return: Лист с кортежами, в котором на первом индексе находится номер страницы, а затем - пять чисел
        1: Координата x штампа
        2: Координата y штампа
        3: Координата x пробивки
        4: Координата y пробивки
        5: Масштаб штампа
        """
        doc_letters = doc_code.split("-")[-1][:3].upper().strip("\n")
        changed_sheets = []
        for change in changes:
            changed_sheets += change["pages"]
        changed_sheets.sort()
        geometry = []
        if doc_letters in DOC_SIZES_MAP.keys():
            coordinates = copy(FORMAT_INFO[DOC_SIZES_MAP[doc_letters]])
            for changed_sheet in changed_sheets:
                geometry.append((changed_sheet, (
                    coordinates["stamp_x"],
                    coordinates["stamp_y"],
                    coordinates["note_x"],
                    coordinates["note_y"],
                    coordinates["scale"]
                )))
        else:
            for changed_sheet in changed_sheets:
                geometry.append((changed_sheet, DEFAULT_GEOMETRY))
        return geometry


if __name__ == "__main__":
    pass
